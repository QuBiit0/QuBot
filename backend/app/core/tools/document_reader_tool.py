"""
Document Reader Tool - Parse PDF, DOCX, XLSX, CSV, and plain text files.
Gives agents the ability to extract content from real-world documents.
"""

import io
import time
from pathlib import Path

from .base import BaseTool, ToolCategory, ToolParameter, ToolResult, ToolRiskLevel


class DocumentReaderTool(BaseTool):
    """
    Read and extract content from PDF, DOCX, XLSX, CSV, and TXT files.
    Supports both local file paths and raw bytes (base64-encoded).
    Use when agents need to process documents, reports, spreadsheets or data files.
    """

    name = "document_reader"
    description = (
        "Extract text and structured data from documents: PDF, DOCX, XLSX, CSV, TXT. "
        "Provide a file path or URL. Returns text content, tables, and metadata. "
        "Use this to read reports, contracts, spreadsheets, or any uploaded file."
    )
    category = ToolCategory.FILE
    risk_level = ToolRiskLevel.SAFE

    MAX_PAGES = 50
    MAX_CHARS = 50_000

    def _get_parameters_schema(self) -> dict[str, ToolParameter]:
        return {
            "path": ToolParameter(
                name="path",
                type="string",
                description="Absolute file path to read (e.g. /tmp/report.pdf)",
                required=True,
            ),
            "pages": ToolParameter(
                name="pages",
                type="string",
                description="Page range for PDFs: '1-5', '3', '1-10'. Default: all (max 50)",
                required=False,
                default=None,
            ),
            "sheet": ToolParameter(
                name="sheet",
                type="string",
                description="Sheet name or index for XLSX files (default: first sheet)",
                required=False,
                default=None,
            ),
            "max_rows": ToolParameter(
                name="max_rows",
                type="integer",
                description="Max rows to return for XLSX/CSV (default 500)",
                required=False,
                default=500,
            ),
            "extract_tables": ToolParameter(
                name="extract_tables",
                type="boolean",
                description="Try to extract tables from PDF/DOCX as structured data",
                required=False,
                default=True,
            ),
        }

    def _validate_config(self) -> None:
        self.allowed_extensions = self.config.get(
            "allowed_extensions", [".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"]
        )

    def _parse_page_range(self, pages_str: str, total_pages: int) -> list[int]:
        """Parse '1-5' or '3' into list of 0-indexed page numbers."""
        result = []
        for part in pages_str.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                result.extend(range(int(start) - 1, min(int(end), total_pages)))
            else:
                idx = int(part) - 1
                if 0 <= idx < total_pages:
                    result.append(idx)
        return sorted(set(result))

    async def _read_pdf(self, path: str, pages: str | None, extract_tables: bool) -> dict:
        try:
            import pdfplumber
        except ImportError:
            # Fallback to PyPDF2
            try:
                import pypdf
                reader = pypdf.PdfReader(path)
                total = len(reader.pages)
                page_nums = (
                    self._parse_page_range(pages, total) if pages
                    else list(range(min(total, self.MAX_PAGES)))
                )
                texts = []
                for i in page_nums:
                    texts.append(f"--- Page {i+1} ---\n{reader.pages[i].extract_text() or ''}")
                return {
                    "text": "\n\n".join(texts),
                    "total_pages": total,
                    "pages_read": len(page_nums),
                    "tables": [],
                }
            except ImportError:
                raise ImportError("Install pdfplumber or pypdf: pip install pdfplumber")

        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
            page_nums = (
                self._parse_page_range(pages, total) if pages
                else list(range(min(total, self.MAX_PAGES)))
            )
            texts = []
            tables = []
            for i in page_nums:
                page = pdf.pages[i]
                text = page.extract_text() or ""
                texts.append(f"--- Page {i+1} ---\n{text}")
                if extract_tables:
                    for table in page.extract_tables():
                        if table:
                            tables.append({"page": i + 1, "data": table})

            return {
                "text": "\n\n".join(texts),
                "total_pages": total,
                "pages_read": len(page_nums),
                "tables": tables[:20],
            }

    async def _read_docx(self, path: str, extract_tables: bool) -> dict:
        try:
            import docx
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        if extract_tables:
            for t in doc.tables:
                rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
                tables.append(rows)

        return {
            "text": "\n\n".join(paragraphs),
            "paragraphs": len(paragraphs),
            "tables": tables[:20],
        }

    async def _read_xlsx(self, path: str, sheet: str | None, max_rows: int) -> dict:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("Install openpyxl: pip install openpyxl")

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        sheet_names = wb.sheetnames

        if sheet is not None:
            ws = wb[sheet] if isinstance(sheet, str) else wb.worksheets[int(sheet)]
        else:
            ws = wb.active

        rows = []
        headers = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c is not None else "" for c in row]
            if i >= max_rows:
                break
            rows.append([str(c) if c is not None else "" for c in row])

        # Build text representation
        lines = []
        for row in rows[:500]:
            lines.append("\t".join(row))
        text = "\n".join(lines)

        return {
            "text": text,
            "headers": headers,
            "rows": rows,
            "total_rows": ws.max_row,
            "total_cols": ws.max_column,
            "sheet_name": ws.title,
            "all_sheets": sheet_names,
        }

    async def _read_csv(self, path: str, max_rows: int) -> dict:
        import csv

        try:
            import chardet
            with open(path, "rb") as f:
                raw = f.read(32768)
            encoding = chardet.detect(raw).get("encoding") or "utf-8"
        except ImportError:
            encoding = "utf-8"

        rows = []
        headers = None
        with open(path, encoding=encoding, errors="replace", newline="") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i == 0:
                    headers = row
                if i >= max_rows:
                    break
                rows.append(row)

        lines = ["\t".join(r) for r in rows[:500]]
        return {
            "text": "\n".join(lines),
            "headers": headers,
            "rows": rows,
            "total_rows": len(rows),
        }

    async def _read_text(self, path: str) -> dict:
        try:
            import chardet
            with open(path, "rb") as f:
                raw = f.read()
            encoding = chardet.detect(raw[:32768]).get("encoding") or "utf-8"
        except ImportError:
            encoding = "utf-8"

        with open(path, encoding=encoding, errors="replace") as f:
            text = f.read(self.MAX_CHARS)

        return {"text": text, "encoding": encoding}

    async def execute(
        self,
        path: str,
        pages: str | None = None,
        sheet: str | None = None,
        max_rows: int = 500,
        extract_tables: bool = True,
    ) -> ToolResult:
        start_time = time.time()

        file_path = Path(path)
        if not file_path.exists():
            return ToolResult(success=False, error=f"File not found: {path}")

        ext = file_path.suffix.lower()
        if ext not in self.allowed_extensions:
            return ToolResult(
                success=False,
                error=f"Unsupported file type '{ext}'. Allowed: {self.allowed_extensions}",
            )

        file_size = file_path.stat().st_size
        if file_size > 100 * 1024 * 1024:  # 100 MB
            return ToolResult(success=False, error="File too large (max 100 MB)")

        try:
            if ext == ".pdf":
                data = await self._read_pdf(path, pages, extract_tables)
            elif ext == ".docx":
                data = await self._read_docx(path, extract_tables)
            elif ext in (".xlsx", ".xls"):
                data = await self._read_xlsx(path, sheet, max_rows)
            elif ext == ".csv":
                data = await self._read_csv(path, max_rows)
            else:
                data = await self._read_text(path)

            # Truncate text
            text = data.get("text", "")
            truncated = False
            if len(text) > self.MAX_CHARS:
                text = text[: self.MAX_CHARS] + "\n\n[Content truncated...]"
                truncated = True
            data["text"] = text

            data.update(
                {
                    "file_name": file_path.name,
                    "file_size": file_size,
                    "extension": ext,
                    "truncated": truncated,
                }
            )

            return ToolResult(
                success=True,
                data=data,
                stdout=text,
                execution_time_ms=int((time.time() - start_time) * 1000),
                metadata={"path": path, "extension": ext, "file_size": file_size},
            )

        except ImportError as e:
            return ToolResult(success=False, error=str(e))
        except Exception as e:
            return ToolResult(
                success=False,
                error=f"Error reading {ext} file: {e}",
                execution_time_ms=int((time.time() - start_time) * 1000),
            )
